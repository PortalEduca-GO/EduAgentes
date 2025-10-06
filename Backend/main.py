import os
import shutil
import threading
from pathlib import Path
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from dotenv import load_dotenv
import uuid

# Carrega as variáveis de ambiente do arquivo .env
load_dotenv()

# Importações dos módulos do projeto
import models
import schemas
import auth
import database

# Importações para o sistema RAG (Retrieval-Augmented Generation)
from langchain_ollama import OllamaLLM as Ollama
from langchain_ollama import OllamaEmbeddings
import importlib

chroma_module_spec = importlib.util.find_spec("langchain_chroma")
if chroma_module_spec is not None:
    langchain_chroma_module = importlib.import_module("langchain_chroma")
    Chroma = getattr(langchain_chroma_module, "Chroma")
    _CHROMA_IMPORT_ERROR: Optional[Exception] = None
else:  # pragma: no cover - depende do ambiente de execução
    Chroma = None  # type: ignore
    _CHROMA_IMPORT_ERROR = ImportError("langchain-chroma não está instalado. Execute 'pip install langchain-chroma'.")

from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import create_retrieval_chain

# Importações para web scraping
import requests
from bs4 import BeautifulSoup
from langchain.schema import Document
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate

# Importação para integração com Google Gemini (importada condicionalmente)
# import google.generativeai as genai  # Movido para dentro do bloco condicional

# Importação para Ollama direto (opcional)
try:
    import ollama
except ImportError:  # pragma: no cover - depende do ambiente
    ollama = None

# Criação das tabelas no banco de dados
models.Base.metadata.create_all(bind=database.engine)

# Configurações do sistema RAG
BASE_DIR = Path(__file__).parent
CHROMA_DB_PATH = str(BASE_DIR / "chroma_db")
UPLOADS_PATH = str(BASE_DIR / "uploads")
STATIC_PATH = str(BASE_DIR / "static")
LOGOS_PATH = f"{STATIC_PATH}/logos"
EMBEDDING_MODEL_NAME = os.getenv("OLLAMA_EMBEDDING_MODEL", "nomic-embed-text")
LLM_MODEL_NAME = os.getenv("OLLAMA_MODEL", "llama3")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
GEMINI_MODEL_NAME = os.getenv("GEMINI_MODEL_NAME", "gemini-2.5-flash")

def extract_smart_content(content_text: str, query: str, max_content: int) -> str:
    """Extrai conteúdo de forma inteligente baseado na query - VERSÃO MENOS RESTRITIVA"""
    
    if len(content_text) <= max_content:
        return content_text
    
    print(f"[SMART_EXTRACT] 🔍 Extraindo conteúdo inteligente para: '{query}'")
    
    # Extrair termos da query de forma mais ampla
    query_terms = query.lower().split()
    
    # Termos específicos ampliados para capturar mais contexto
    specific_terms = [
        "agrocolégio", "agrocolegio", "maguito", "vilela",
        "escola", "educação", "ensino", "estadual", "unidade",
        "estudante", "aluno", "professor", "diretor"
    ]
    
    all_terms = list(set(query_terms + specific_terms))  # Remove duplicatas
    print(f"[SMART_EXTRACT] 📋 Termos de busca: {all_terms}")
    
    # Início maior do documento para mais contexto
    start_content = content_text[:5000]  # Aumentado de 2000 para 5000
    
    # Buscar seções relevantes com critérios mais amplos
    remaining_content = content_text[5000:]
    found_sections = []
    
    for term in all_terms:
        if len(term) > 2:  # Reduzido de 3 para 2 (menos restritivo)
            term_lower = term.lower()
            content_lower = remaining_content.lower()
            
            start_pos = 0
            sections_for_term = 0
            while True:
                pos = content_lower.find(term_lower, start_pos)
                if pos == -1 or sections_for_term >= 4:  # Aumentado de 2 para 4 seções por termo
                    break
                
                # Extrair contexto MAIOR (800 chars antes e depois)
                context_start = max(0, pos - 800)  # Aumentado de 400 para 800
                context_end = min(len(remaining_content), pos + 800)  # Aumentado de 400 para 800
                context_section = remaining_content[context_start:context_end]
                
                # Aceitar mais seções (até 15 em vez de 8)
                if context_section not in found_sections and len(found_sections) < 15:
                    found_sections.append(context_section)
                    sections_for_term += 1
                    print(f"[SMART_EXTRACT] ✅ Seção encontrada para '{term}': {len(context_section)} chars")
                
                start_pos = pos + 1
    
    # Combinar mais seções
    if found_sections:
        relevant_sections = found_sections[:8]  # Aumentado de 4 para 8 seções
        relevant_content = "\n\n[...TRECHO RELEVANTE...]\n\n".join(relevant_sections)
        limited_content = start_content + "\n\n[...CONTEÚDO RELEVANTE PARA A PERGUNTA...]\n\n" + relevant_content
        print(f"[SMART_EXTRACT] ✅ {len(found_sections)} seções relevantes encontradas")
    else:
        # Fallback: mais conteúdo do início
        limited_content = content_text[:max_content * 2]  # Dobrar o limite se não encontrar nada específico
        print(f"[SMART_EXTRACT] ⚠️ Nenhuma seção específica encontrada, usando início estendido")
    
    result = limited_content + "\n\n[Documento otimizado para a pergunta - versão expandida]"
    print(f"[SMART_EXTRACT] 📊 Resultado final: {len(result)} caracteres")
    return result

# Criar diretórios necessários
for path in [CHROMA_DB_PATH, UPLOADS_PATH, STATIC_PATH, LOGOS_PATH]:
    os.makedirs(path, exist_ok=True)

# Inicialização preguiçosa dos componentes de IA
embeddings: Optional[OllamaEmbeddings] = None
llm: Optional[Ollama] = None
_embeddings_lock = threading.Lock()
_llm_lock = threading.Lock()


def _ollama_available() -> bool:
    return ollama is not None


def get_embeddings() -> Optional[OllamaEmbeddings]:
    """Obtém (ou inicializa) a instância de embeddings do Ollama."""
    global embeddings
    if embeddings is not None:
        return embeddings

    if not _ollama_available():
        print("[OLLAMA] Cliente Python não disponível. Instale o pacote 'ollama'.")
        return None

    with _embeddings_lock:
        if embeddings is None:
            try:
                embeddings = OllamaEmbeddings(
                    model=EMBEDDING_MODEL_NAME,
                    base_url=OLLAMA_BASE_URL
                )
            except Exception as exc:
                print(f"[OLLAMA] Falha ao inicializar embeddings: {exc}")
                embeddings = None
    return embeddings


def get_llm() -> Optional[Ollama]:
    """Obtém (ou inicializa) a instância LLM do Ollama."""
    global llm
    if llm is not None:
        return llm

    if not _ollama_available():
        print("[OLLAMA] Cliente Python não disponível. Instale o pacote 'ollama'.")
        return None

    with _llm_lock:
        if llm is None:
            try:
                llm = Ollama(
                    model=LLM_MODEL_NAME,
                    base_url=OLLAMA_BASE_URL,
                    timeout=int(os.getenv("OLLAMA_TIMEOUT_SECONDS", "300")),
                    verbose=os.getenv("OLLAMA_VERBOSE", "false").lower() in {"1", "true", "yes"}
                )
            except Exception as exc:
                print(f"[OLLAMA] Falha ao inicializar LLM: {exc}")
                llm = None
    return llm


def build_chroma(persist_directory: str, embedding_function):
    """Cria instância do Chroma garantindo dependências disponíveis."""
    if Chroma is None:
        message = str(_CHROMA_IMPORT_ERROR) if _CHROMA_IMPORT_ERROR else "Dependência langchain-chroma ausente."
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=message
        )
    return Chroma(persist_directory=persist_directory, embedding_function=embedding_function)


def get_ollama_client():
    """Retorna cliente Ollama configurado, se disponível."""
    if not _ollama_available():
        return None
    try:
        return ollama.Client(host=OLLAMA_BASE_URL)
    except AttributeError:
        # Versões antigas utilizam parâmetro base_url diretamente
        return None
    except Exception as exc:
        print(f"[OLLAMA] Falha ao criar cliente: {exc}")
        return None

# Configuração do Google Gemini
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if GOOGLE_API_KEY:
    import google.generativeai as genai
    genai.configure(api_key=GOOGLE_API_KEY)

# Inicialização da aplicação FastAPI
app = FastAPI(
    title="Edu API",
    description="API para a plataforma Edu com agentes de IA híbridos e sistema RAG em três estágios.",
    version="2.0.0",
    redirect_slashes=False,
)

# Configuração do CORS para permitir comunicação com o Frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000", 
        "http://localhost:3001", 
        "http://127.0.0.1:3000", 
        "http://127.0.0.1:3001",
        "https://teamsbot.dev.educacao.go.gov.br"  # Domínio de produção
    ],  # Frontend URLs
    allow_credentials=True,
    allow_methods=["*"],  # Permite todos os métodos HTTP
    allow_headers=["*"],  # Permite todos os headers
)

# Servir arquivos estáticos (logos dos agentes)
app.mount("/static", StaticFiles(directory=STATIC_PATH), name="static")

# === ENDPOINT RAIZ ===
@app.get("/")
async def root():
    """Endpoint raiz"""
    return {"message": "API Edu", "status": "online"}

# === ENDPOINT DE HEALTH CHECK ===
@app.get("/health")
async def health_check():
    """Endpoint para verificar se a API está funcionando"""
    return {
        "status": "online",
        "message": "API Edu está funcionando",
        "version": "2.0.0"
    }

# === ENDPOINTS DE AUTENTICAÇÃO E USUÁRIOS ===

@app.post("/register", response_model=schemas.User, status_code=status.HTTP_201_CREATED)
def register_user(
    user: schemas.UserCreate, 
    db: Session = Depends(database.get_db)
) -> schemas.User:
    """
    Registra um novo usuário no sistema.
    
    Args:
        user: Dados do usuário a ser criado
        db: Sessão do banco de dados
        
    Returns:
        schemas.User: Dados do usuário criado
        
    Raises:
        HTTPException: Se o nome de usuário já estiver em uso
    """
    db_user = db.query(models.User).filter(models.User.username == user.username).first()
    if db_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Username already registered"
        )
    
    hashed_password = auth.get_password_hash(user.password)
    new_user = models.User(
        username=user.username, 
        hashed_password=hashed_password,
        role=models.UserRole.USER
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return new_user

@app.post("/login", response_model=schemas.Token)
def login_for_access_token(
    form_data: OAuth2PasswordRequestForm = Depends(), 
    db: Session = Depends(database.get_db)
) -> schemas.Token:
    """
    Autentica um usuário e retorna um token de acesso.
    
    Args:
        form_data: Dados do formulário de login
        db: Sessão do banco de dados
        
    Returns:
        schemas.Token: Token de acesso JWT
        
    Raises:
        HTTPException: Se as credenciais estiverem incorretas
    """
    user = db.query(models.User).filter(models.User.username == form_data.username).first()
    if not user or not auth.verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token = auth.create_access_token(data={"sub": user.username})
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/users/me", response_model=schemas.User)
def read_users_me(current_user: models.User = Depends(auth.get_current_user)) -> schemas.User:
    """
    Retorna os dados do usuário autenticado.
    
    Args:
        current_user: Usuário atual autenticado
        
    Returns:
        schemas.User: Dados do usuário
    """
    return current_user

# === ENDPOINTS DE GERENCIAMENTO DE AGENTES ===

@app.post("/agents", response_model=schemas.Agent, status_code=status.HTTP_201_CREATED)
def create_agent(
    agent: schemas.AgentCreate, 
    db: Session = Depends(database.get_db), 
    current_user: models.User = Depends(auth.get_current_user)
) -> schemas.Agent:
    """
    Cria um novo agente para o usuário autenticado.
    
    Args:
        agent: Dados do agente a ser criado
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        schemas.Agent: Dados do agente criado
    """
    new_agent = models.Agent(
        name=agent.name,
        description=agent.description,
        system_prompt=agent.system_prompt,
        owner_id=current_user.id
    )
    db.add(new_agent)
    db.commit()
    db.refresh(new_agent)
    return new_agent

@app.get("/agents", response_model=List[schemas.Agent])
def read_agents_list(
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> List[schemas.Agent]:
    """
    Lista todos os agentes aprovados disponíveis para uso.
    
    Args:
        db: Sessão do banco de dados
        
    Returns:
        List[schemas.Agent]: Lista de agentes aprovados
    """
    agents = db.query(models.Agent).filter(
        models.Agent.status == models.AgentStatus.APPROVED
    ).all()
    return agents

@app.get("/agents/pending", response_model=List[schemas.Agent])
def get_pending_agents(
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))  # Temporariamente ADMIN
) -> List[schemas.Agent]:
    """
    [MASTER_ADMIN] Lista todos os agentes pendentes de aprovação.
    
    Args:
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões adequadas
        
    Returns:
        List[schemas.Agent]: Lista de agentes pendentes
    """
    pending_agents = db.query(models.Agent).filter(
        models.Agent.status == models.AgentStatus.PENDING
    ).all()
    return pending_agents

@app.get("/agents/{agent_id}", response_model=schemas.Agent)
def read_agent_details(
    agent_id: int, 
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> schemas.Agent:
    """
    Busca os detalhes de um agente específico.
    
    Args:
        agent_id: ID do agente
        db: Sessão do banco de dados
        
    Returns:
        schemas.Agent: Dados do agente
        
    Raises:
        HTTPException: Se o agente não for encontrado
    """
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if db_agent is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    return db_agent

# === ENDPOINTS DE ADMINISTRAÇÃO ===

@app.patch("/agents/{agent_id}/status", response_model=schemas.Agent)
def update_agent_status(
    agent_id: int,
    status_update: schemas.AgentStatusUpdate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> schemas.Agent:
    """
    [MASTER_ADMIN] Atualiza o status de aprovação de um agente.
    
    Args:
        agent_id: ID do agente
        status_update: Novo status do agente
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões adequadas
        
    Returns:
        schemas.Agent: Dados do agente atualizado
        
    Raises:
        HTTPException: Se o agente não for encontrado
    """
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    db_agent.status = status_update.status
    db.commit()
    db.refresh(db_agent)
    return db_agent

# ==================== NOVAS ROTAS DE GERENCIAMENTO ====================

@app.get("/admin/agents", response_model=List[schemas.Agent])
def get_all_agents_admin(
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
) -> List[schemas.Agent]:
    """
    [ADMIN+] Lista TODOS os agentes do sistema (independente do status).
    
    Args:
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Admin
        
    Returns:
        List[schemas.Agent]: Lista completa de agentes
    """
    agents = db.query(models.Agent).all()
    return agents

@app.patch("/agents/{agent_id}", response_model=schemas.Agent)
def update_agent(
    agent_id: int,
    agent_update: schemas.AgentUpdate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
) -> schemas.Agent:
    """
    [ADMIN+] Atualiza um agente existente.
    
    Args:
        agent_id: ID do agente a ser atualizado
        agent_update: Dados para atualização
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Admin
        
    Returns:
        schemas.Agent: Agente atualizado
        
    Raises:
        HTTPException: Se o agente não for encontrado
    """
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    # Atualiza apenas os campos fornecidos
    update_data = agent_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_agent, field, value)
    
    db.commit()
    db.refresh(db_agent)
    return db_agent

@app.put("/agents/{agent_id}", response_model=schemas.Agent)
def update_agent(
    agent_id: int,
    agent_update: schemas.AgentUpdate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> schemas.Agent:
    """
    [MASTER_ADMIN] Atualiza um agente existente.
    
    Args:
        agent_id: ID do agente a ser atualizado
        agent_update: Dados de atualização do agente
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Master Admin
        
    Returns:
        schemas.Agent: Agente atualizado
        
    Raises:
        HTTPException: Se o agente não for encontrado
    """
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    # Verifica se é o proprietário ou Master Admin
    if current_user.role != models.UserRole.MASTER_ADMIN and db_agent.owner_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Não autorizado a editar este agente"
        )
    
    # Atualiza apenas os campos fornecidos
    update_data = agent_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_agent, field, value)
    
    print(f"[ADMIN] Agente {agent_id} atualizado por {current_user.username}")
    db.commit()
    db.refresh(db_agent)
    return db_agent

@app.delete("/agents/{agent_id}")
def delete_agent(
    agent_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
):
    """
    [MASTER_ADMIN] Exclui um agente e toda sua base de conhecimento.
    
    Args:
        agent_id: ID do agente a ser excluído
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Master Admin
        
    Returns:
        dict: Mensagem de confirmação
        
    Raises:
        HTTPException: Se o agente não for encontrado
    """
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    try:
        # Remove arquivos de documentos
        documents = db.query(models.Document).filter(models.Document.agent_id == agent_id).all()
        for doc in documents:
            if os.path.exists(doc.file_path):
                os.remove(doc.file_path)
        
        # Remove pasta de vetores do Chroma DB
        agent_chroma_path = os.path.join(CHROMA_DB_PATH, str(agent_id))
        if os.path.exists(agent_chroma_path):
            shutil.rmtree(agent_chroma_path)
        
        # Remove logo se existir
        if db_agent.logo_url:
            logo_path = os.path.join(".", db_agent.logo_url.lstrip("/"))
            if os.path.exists(logo_path):
                os.remove(logo_path)
        
        # O SQLAlchemy irá remover os documentos e links automaticamente
        # devido ao cascade="all, delete-orphan" definido no modelo
        db.delete(db_agent)
        db.commit()
        
        return {"message": f"Agent {agent_id} and all associated data deleted successfully"}
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting agent: {str(e)}"
        )

@app.get("/agents/{agent_id}/documents", response_model=List[schemas.Document])
def get_agent_documents(
    agent_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
) -> List[schemas.Document]:
    """
    [ADMIN+] Lista todos os documentos de um agente.
    
    Args:
        agent_id: ID do agente
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        List[schemas.Document]: Lista de documentos do agente
    """
    # Verifica se o agente existe
    agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    documents = db.query(models.Document).filter(models.Document.agent_id == agent_id).all()
    return documents

@app.delete("/documents/{document_id}")
def delete_document(
    document_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
):
    """
    [ADMIN+] Exclui um documento específico.
    
    Args:
        document_id: ID do documento a ser excluído
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        dict: Mensagem de confirmação
    """
    document = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        # Remove arquivo físico
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
        
        # Remove do banco de dados
        db.delete(document)
        db.commit()
        
        return {"message": f"Document {document_id} deleted successfully"}
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting document: {str(e)}"
        )

@app.delete("/links/{link_id}")
def delete_link(
    link_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
):
    """
    [ADMIN+] Exclui um link específico.
    
    Args:
        link_id: ID do link a ser excluído
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        dict: Mensagem de confirmação
    """
    link = db.query(models.Link).filter(models.Link.id == link_id).first()
    if not link:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Link not found"
        )
    
    try:
        db.delete(link)
        db.commit()
        return {"message": f"Link {link_id} deleted successfully"}
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting link: {str(e)}"
        )

@app.get("/master/users", response_model=List[schemas.User])
def get_all_users(
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> List[schemas.User]:
    """
    [MASTER_ADMIN] Lista todos os usuários do sistema.
    
    Args:
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Master Admin
        
    Returns:
        List[schemas.User]: Lista de todos os usuários
    """
    users = db.query(models.User).all()
    return users

@app.post("/master/users", response_model=schemas.User, status_code=status.HTTP_201_CREATED)
def create_user_by_master(
    user: schemas.UserCreateByMaster,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> schemas.User:
    """
    [MASTER_ADMIN] Cria um novo usuário com papel específico.
    
    Args:
        user: Dados do usuário a ser criado
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Master Admin
        
    Returns:
        schemas.User: Dados do usuário criado
        
    Raises:
        HTTPException: Se o nome de usuário já estiver em uso
    """
    print(f"DEBUG CREATE USER: Dados recebidos: {user}")
    print(f"DEBUG CREATE USER: Username: '{user.username}', Password length: {len(user.password)}, Role: {user.role}")
    
    db_user = db.query(models.User).filter(models.User.username == user.username).first()
    if db_user:
        print(f"DEBUG CREATE USER: Usuário '{user.username}' já existe!")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Username already registered"
        )
    
    hashed_password = auth.get_password_hash(user.password)
    new_user = models.User(
        username=user.username, 
        hashed_password=hashed_password,
        role=user.role
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    print(f"DEBUG CREATE USER: Usuário '{user.username}' criado com sucesso!")
    return new_user

# ==================== ENDPOINTS DE CONFIGURAÇÃO DO SISTEMA ====================

@app.get("/master/system/ai-model", response_model=schemas.SystemConfigResponse)
def get_ai_model_config(
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> schemas.SystemConfigResponse:
    """
    [MASTER_ADMIN] Obtém a configuração atual do modelo de IA do sistema.
    
    Returns:
        SystemConfigResponse: Configuração atual do modelo de IA
    """
    config = db.query(models.SystemConfig).filter(
        models.SystemConfig.key == "ai_model_type"
    ).first()
    
    if not config:
        # Cria configuração padrão se não existir
        config = models.SystemConfig(
            key="ai_model_type",
            value="HYBRID",
            description="Tipo de modelo de IA usado no sistema (HYBRID, LLAMA_ONLY, GEMINI_ONLY)",
            updated_by=current_user.id
        )
        db.add(config)
        db.commit()
        db.refresh(config)
    
    # Busca o usuário que fez a última atualização
    updated_by_user = db.query(models.User).filter(models.User.id == config.updated_by).first()
    
    return schemas.SystemConfigResponse(
        key=config.key,
        value=config.value,
        description=config.description,
        updated_at=config.updated_at,
        updated_by_username=updated_by_user.username if updated_by_user else "Sistema"
    )

@app.put("/master/system/ai-model", response_model=schemas.SystemConfigResponse)
def update_ai_model_config(
    config_update: schemas.AIModelConfigUpdate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> schemas.SystemConfigResponse:
    """
    [MASTER_ADMIN] Atualiza a configuração do modelo de IA do sistema.
    
    Args:
        config_update: Nova configuração do modelo de IA
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Master Admin
        
    Returns:
        SystemConfigResponse: Configuração atualizada
    """
    # Valida o tipo de modelo
    valid_models = ["HYBRID", "LLAMA_ONLY", "GEMINI_ONLY"]
    if config_update.ai_model_type not in valid_models:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Modelo de IA inválido. Opções válidas: {', '.join(valid_models)}"
        )
    
    # Busca ou cria a configuração
    config = db.query(models.SystemConfig).filter(
        models.SystemConfig.key == "ai_model_type"
    ).first()
    
    if config:
        config.value = config_update.ai_model_type
        config.updated_by = current_user.id
    else:
        config = models.SystemConfig(
            key="ai_model_type",
            value=config_update.ai_model_type,
            description="Tipo de modelo de IA usado no sistema (HYBRID, LLAMA_ONLY, GEMINI_ONLY)",
            updated_by=current_user.id
        )
        db.add(config)
    
    db.commit()
    db.refresh(config)
    
    # Busca o usuário que fez a atualização
    updated_by_user = db.query(models.User).filter(models.User.id == config.updated_by).first()
    
    print(f"[SISTEMA] Modelo de IA alterado para: {config_update.ai_model_type} por {current_user.username}")
    
    return schemas.SystemConfigResponse(
        key=config.key,
        value=config.value,
        description=config.description,
        updated_at=config.updated_at,
        updated_by_username=updated_by_user.username if updated_by_user else "Sistema"
    )

# === ENDPOINTS DE PERSONALIZAÇÃO DE AGENTES ===

@app.post("/agents/{agent_id}/logo", response_model=schemas.LogoUploadResponse)
async def upload_agent_logo(
    agent_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> schemas.LogoUploadResponse:
    """
    Upload de logo para um agente específico.
    
    Args:
        agent_id: ID do agente
        file: Arquivo de imagem do logo
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        schemas.LogoUploadResponse: Status do upload e URL do logo
        
    Raises:
        HTTPException: Se o agente não for encontrado ou arquivo não for válido
    """
    print(f"DEBUG LOGO UPLOAD: agent_id={agent_id}, filename={file.filename}, content_type={file.content_type}")
    print(f"DEBUG LOGO UPLOAD: file size={file.size if hasattr(file, 'size') else 'unknown'}")
    
    # Verifica se o agente existe e se o usuário tem permissão
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        print(f"DEBUG LOGO UPLOAD: Agent {agent_id} not found")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    # Verifica se o usuário é o dono do agente ou é admin
    if (db_agent.owner_id != current_user.id and 
        current_user.role not in [models.UserRole.ADMIN, models.UserRole.MASTER_ADMIN]):
        print(f"DEBUG LOGO UPLOAD: User {current_user.username} doesn't have permission for agent {agent_id}")
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions"
        )
    
    print(f"DEBUG LOGO UPLOAD: Permission check passed for user {current_user.username}")
    
    # Verifica se o arquivo é uma imagem
    allowed_types = ["image/jpeg", "image/png", "image/gif", "image/webp", "image/jpg"]
    print(f"DEBUG LOGO UPLOAD: Checking content type '{file.content_type}' against allowed types {allowed_types}")
    if file.content_type not in allowed_types:
        print(f"DEBUG LOGO UPLOAD: Invalid content type: {file.content_type}")
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Apenas arquivos de imagem são permitidos (JPEG, PNG, GIF, WebP). Tipo recebido: {file.content_type}"
        )
    
    # Verifica se o arquivo tem nome e extensão
    print(f"DEBUG LOGO UPLOAD: Checking filename: '{file.filename}'")
    if not file.filename:
        print("DEBUG LOGO UPLOAD: No filename provided")
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Nome do arquivo é obrigatório"
        )
    
    if '.' not in file.filename:
        print(f"DEBUG LOGO UPLOAD: No extension in filename: '{file.filename}'")
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Arquivo deve ter uma extensão válida"
        )
    
    try:
        # Gera nome único para o arquivo
        file_extension = file.filename.split('.')[-1].lower()
        if file_extension not in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Extensão de arquivo não suportada: {file_extension}"
            )
            
        unique_filename = f"agent_{agent_id}_{uuid.uuid4().hex}.{file_extension}"
        logo_path = f"{LOGOS_PATH}/{unique_filename}"
        
        # Salva o arquivo
        with open(logo_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Atualiza o campo logo_url no banco de dados
        logo_url = f"/static/logos/{unique_filename}"
        db_agent.logo_url = logo_url
        db.commit()
        db.refresh(db_agent)
        
        return schemas.LogoUploadResponse(
            status="success",
            message=f"Logo carregado com sucesso para o agente '{db_agent.name}'",
            logo_url=logo_url
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao processar o logo: {str(e)}"
        )

# === ENDPOINTS DE GERENCIAMENTO DE LINKS ===

@app.post("/agents/{agent_id}/links", response_model=schemas.LinkCreateResponse)
def add_link_to_agent(
    agent_id: int,
    link_data: schemas.LinkCreate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> schemas.LinkCreateResponse:
    """
    Adiciona um novo link à base de conhecimento do agente.
    
    Args:
        agent_id: ID do agente
        link_data: Dados do link a ser adicionado
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        schemas.LinkCreateResponse: Status da operação e ID do link criado
        
    Raises:
        HTTPException: Se o agente não for encontrado
    """
    # Verifica se o agente existe e se o usuário tem permissão
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    # Verifica se o usuário é o dono do agente ou é admin
    if (db_agent.owner_id != current_user.id and 
        current_user.role not in [models.UserRole.ADMIN, models.UserRole.MASTER_ADMIN]):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions"
        )
    
    try:
        # Cria o registro do link no banco de dados
        new_link = models.Link(
            url=link_data.url,
            title=link_data.title,
            description=link_data.description,
            agent_id=agent_id
        )

        db.add(new_link)
        db.commit()
        db.refresh(new_link)

        # Tenta automaticamente fazer scraping e adicionar ao Chroma
        # (Não é fatal: se o scraping falhar, o link continua registrado como metadado)
        try:
            print(f"[LINK INGEST] Tentando fazer scraping automático do link: {new_link.url}")
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            resp = requests.get(new_link.url, headers=headers, timeout=15)
            resp.raise_for_status()

            soup = BeautifulSoup(resp.content, 'html.parser')
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()

            title = soup.find('title')
            title_text = title.get_text().strip() if title else (new_link.title or "Sem título")

            text_content = soup.get_text()
            lines = [line.strip() for line in text_content.splitlines()]
            content = '\n'.join(line for line in lines if line)

            if content and len(content.strip()) >= 100:
                # Cria documento e divide em chunks
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": new_link.url,
                        "title": title_text,
                        "type": "web_scraping",
                        "agent_id": agent_id
                    }
                )

                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=2000,  # Aumentado para capturar mais contexto
                    chunk_overlap=400  # Aumentado proporcionalmente
                )
                chunks = text_splitter.split_documents([doc])

                # Adiciona ao Chroma do agente (usa mesmo padrão que upload)
                agent_chroma_path = f"{CHROMA_DB_PATH}/{agent_id}"
                os.makedirs(agent_chroma_path, exist_ok=True)
                embedding_function = get_embeddings()
                if embedding_function is None:
                    print("[LINK INGEST] Embeddings indisponíveis - verifique o serviço Ollama")
                else:
                    agent_vectorstore = build_chroma(
                        persist_directory=agent_chroma_path,
                        embedding_function=embedding_function
                    )
                try:
                    if embedding_function is not None:
                        agent_vectorstore.add_documents(documents=chunks)
                        print(f"[LINK INGEST] Conteúdo do link adicionado ao Chroma ({len(chunks)} chunks)")
                except Exception as e:
                    print(f"[LINK INGEST] Falha ao adicionar documentos ao Chroma: {e}")
            else:
                print(f"[LINK INGEST] Conteúdo insuficiente extraído de {new_link.url}")

        except Exception as e:
            # Loga erro mas não impede o sucesso do endpoint de criação do link
            print(f"[LINK INGEST] Erro ao extrair conteúdo do link {new_link.url}: {e}")

        return schemas.LinkCreateResponse(
            status="success",
            message=f"Link adicionado com sucesso ao agente '{db_agent.name}'",
            id=new_link.id,
            title=new_link.title or link_data.url
        )

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao adicionar link: {str(e)}"
        )

@app.get("/agents/{agent_id}/links", response_model=List[schemas.Link])
def get_agent_links(
    agent_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> List[schemas.Link]:
    """
    Lista todos os links de um agente.
    
    Args:
        agent_id: ID do agente
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        List[schemas.Link]: Lista de links do agente
        
    Raises:
        HTTPException: Se o agente não for encontrado
    """
    # Verifica se o agente existe
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    # Busca os links do agente
    links = db.query(models.Link).filter(models.Link.agent_id == agent_id).all()
    return links

@app.delete("/links/{link_id}")
def delete_link(
    link_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> Dict[str, str]:
    """
    Remove um link específico.
    
    Args:
        link_id: ID do link a ser removido
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        Dict: Status da operação
        
    Raises:
        HTTPException: Se o link não for encontrado ou usuário não tiver permissão
    """
    # Busca o link
    db_link = db.query(models.Link).filter(models.Link.id == link_id).first()
    if not db_link:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Link not found"
        )
    
    # Verifica se o usuário tem permissão
    db_agent = db.query(models.Agent).filter(models.Agent.id == db_link.agent_id).first()
    if (db_agent.owner_id != current_user.id and 
        current_user.role not in [models.UserRole.ADMIN, models.UserRole.MASTER_ADMIN]):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions"
        )
    
    # Remove o link
    db.delete(db_link)
    db.commit()
    
    return {"status": "success", "message": "Link removido com sucesso"}

# === ENDPOINTS DA BASE DE CONHECIMENTO CENTRALIZADA ===

@app.get("/knowledge", response_model=List[schemas.Knowledge])
def get_knowledge_list(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> List[schemas.Knowledge]:
    """
    Lista todos os conhecimentos da base centralizada.
    
    Args:
        skip: Número de itens para pular
        limit: Limite de itens a retornar
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        List[schemas.Knowledge]: Lista de conhecimentos
    """
    try:
        print(f"DEBUG: Getting knowledge list for user {current_user.username}")
        knowledge_items = db.query(models.Knowledge).filter(
            models.Knowledge.status == models.KnowledgeStatus.APPROVED
        ).order_by(models.Knowledge.id).offset(skip).limit(limit).all()
        print(f"DEBUG: Found {len(knowledge_items)} knowledge items")
        
        result = []
        for item in knowledge_items:
            try:
                schema_item = schemas.Knowledge.from_orm(item)
                result.append(schema_item)
                print(f"DEBUG: Successfully converted item {item.id}")
            except Exception as e:
                print(f"DEBUG: Error converting item {item.id}: {e}")
                import traceback
                traceback.print_exc()
                raise
        
        print(f"DEBUG: Returning {len(result)} items")
        return result
    except Exception as e:
        print(f"DEBUG: Error in get_knowledge_list: {e}")
        import traceback
        traceback.print_exc()
        raise
    return knowledge_items

@app.get("/knowledge/pending", response_model=List[schemas.Knowledge])
def get_pending_knowledge(
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> List[schemas.Knowledge]:
    """
    [MASTER_ADMIN] Lista conhecimentos pendentes de aprovação.
    
    Args:
        db: Sessão do banco de dados
        current_user: Usuário Master Admin autenticado
        
    Returns:
        List[schemas.Knowledge]: Lista de conhecimentos pendentes
    """
    from sqlalchemy.orm import joinedload
    
    try:
        print(f"Buscando conhecimentos pendentes para usuário: {current_user.username}")
        
        # Buscar conhecimentos com status PENDING
        pending_knowledge = db.query(models.Knowledge).options(
            joinedload(models.Knowledge.author),
            joinedload(models.Knowledge.approved_by),
            joinedload(models.Knowledge.agents)
        ).filter(
            models.Knowledge.status == models.KnowledgeStatus.PENDING
        ).order_by(models.Knowledge.created_at.desc()).all()
        
        print(f"Encontrados {len(pending_knowledge)} conhecimentos pendentes")
        
        return pending_knowledge
        
    except Exception as e:
        print(f"Erro ao buscar conhecimentos pendentes: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro interno do servidor: {str(e)}"
        )

@app.get("/knowledge/{knowledge_id}", response_model=schemas.Knowledge)
def get_knowledge_detail(
    knowledge_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> schemas.Knowledge:
    """
    Obtém detalhes de um conhecimento específico.
    
    Args:
        knowledge_id: ID do conhecimento
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        schemas.Knowledge: Dados do conhecimento
        
    Raises:
        HTTPException: Se o conhecimento não for encontrado
    """
    knowledge = db.query(models.Knowledge).filter(models.Knowledge.id == knowledge_id).first()
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Knowledge not found"
        )
    return knowledge

@app.post("/knowledge", response_model=schemas.Knowledge, status_code=status.HTTP_201_CREATED)
def create_knowledge(
    knowledge_data: schemas.KnowledgeCreate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
) -> schemas.Knowledge:
    """
    [ADMIN+] Cria um novo item de conhecimento.
    
    Args:
        knowledge_data: Dados do conhecimento a ser criado
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões adequadas
        
    Returns:
        schemas.Knowledge: Conhecimento criado
    """
    try:
        # Define status baseado no papel do usuário
        initial_status = models.KnowledgeStatus.APPROVED if current_user.role == models.UserRole.MASTER_ADMIN else models.KnowledgeStatus.PENDING
        
        # Cria o item de conhecimento
        new_knowledge = models.Knowledge(
            title=knowledge_data.title,
            content=knowledge_data.content,
            knowledge_type=knowledge_data.knowledge_type,
            url=knowledge_data.url,
            tags=knowledge_data.tags,
            expires_at=knowledge_data.expires_at,
            author_id=current_user.id,
            status=initial_status
        )
        
        # Se o Master Admin criou, já marca como aprovado
        if current_user.role == models.UserRole.MASTER_ADMIN:
            from datetime import datetime
            new_knowledge.approved_at = datetime.now()
            new_knowledge.approved_by_id = current_user.id
        
        db.add(new_knowledge)
        db.flush()  # Para obter o ID
        
        # Associa com agentes se especificado
        if knowledge_data.agent_ids:
            agents = db.query(models.Agent).filter(
                models.Agent.id.in_(knowledge_data.agent_ids)
            ).all()
            new_knowledge.agents = agents
        
        db.commit()
        db.refresh(new_knowledge)
        
        return new_knowledge
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao criar conhecimento: {str(e)}"
        )

@app.put("/knowledge/{knowledge_id}", response_model=schemas.Knowledge)
def update_knowledge(
    knowledge_id: int,
    knowledge_update: schemas.KnowledgeUpdate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
) -> schemas.Knowledge:
    """
    [ADMIN+] Atualiza um conhecimento existente.
    
    Args:
        knowledge_id: ID do conhecimento a ser atualizado
        knowledge_update: Dados para atualização
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        schemas.Knowledge: Conhecimento atualizado
    """
    knowledge = db.query(models.Knowledge).filter(models.Knowledge.id == knowledge_id).first()
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Knowledge not found"
        )
    
    # Verifica se é o autor ou admin superior
    if (knowledge.author_id != current_user.id and 
        current_user.role not in [models.UserRole.MASTER_ADMIN, models.UserRole.MASTER_ADMIN]):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions"
        )
    
    try:
        # Atualiza campos fornecidos
        update_data = knowledge_update.dict(exclude_unset=True, exclude={'agent_ids'})
        for field, value in update_data.items():
            setattr(knowledge, field, value)
        
        # Atualiza associações com agentes se fornecido
        if knowledge_update.agent_ids is not None:
            agents = db.query(models.Agent).filter(
                models.Agent.id.in_(knowledge_update.agent_ids)
            ).all()
            knowledge.agents = agents
        
        db.commit()
        db.refresh(knowledge)
        
        return knowledge
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao atualizar conhecimento: {str(e)}"
        )

@app.delete("/knowledge/{knowledge_id}")
def delete_knowledge(
    knowledge_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
):
    """
    [ADMIN+] Exclui um conhecimento.
    
    Args:
        knowledge_id: ID do conhecimento a ser excluído
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        dict: Mensagem de confirmação
    """
    knowledge = db.query(models.Knowledge).filter(models.Knowledge.id == knowledge_id).first()
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Knowledge not found"
        )
    
    # Verifica se é o autor ou admin superior
    if (knowledge.author_id != current_user.id and 
        current_user.role not in [models.UserRole.MASTER_ADMIN, models.UserRole.MASTER_ADMIN]):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions"
        )
    
    try:
        # Remove arquivo se existir
        if knowledge.file_path and os.path.exists(knowledge.file_path):
            os.remove(knowledge.file_path)
        
        db.delete(knowledge)
        db.commit()
        
        return {"message": f"Conhecimento '{knowledge.title}' excluído com sucesso"}
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao excluir conhecimento: {str(e)}"
        )

@app.post("/knowledge/upload", response_model=schemas.KnowledgeUploadResponse)
async def upload_knowledge_document(
    title: str = Form(...),
    agent_ids: str = Form(""),  # IDs separados por vírgula
    tags: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.ADMIN))
) -> schemas.KnowledgeUploadResponse:
    """
    [ADMIN+] Upload de documento para a base de conhecimento.
    
    Args:
        title: Título do conhecimento
        agent_ids: IDs dos agentes separados por vírgula
        tags: Tags separadas por vírgula
        file: Arquivo a ser carregado
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        schemas.KnowledgeUploadResponse: Status do upload
    """
    print(f"DEBUG Upload - title: {title}")
    print(f"DEBUG Upload - agent_ids: {agent_ids}")
    print(f"DEBUG Upload - tags: {tags}")
    print(f"DEBUG Upload - file: {file.filename}, content_type: {file.content_type}")
    
    # Verifica tipo de arquivo
    allowed_types = [
        "application/pdf", 
        "text/plain", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ]
    print(f"DEBUG Upload - file content_type: {file.content_type}")
    print(f"DEBUG Upload - allowed types: {allowed_types}")
    
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Tipo de arquivo não permitido: {file.content_type}. Apenas arquivos PDF, DOCX e TXT são permitidos."
        )
    
    # Salva arquivo
    knowledge_uploads_path = f"{UPLOADS_PATH}/knowledge"
    os.makedirs(knowledge_uploads_path, exist_ok=True)
    
    file_path = f"{knowledge_uploads_path}/{file.filename}"
    
    try:
        content = await file.read()
        
        # Verifica tamanho do arquivo (limite: 5MB)
        MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB em bytes
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"Arquivo muito grande. Tamanho máximo permitido: 5MB. Tamanho atual: {len(content) / (1024*1024):.1f}MB"
            )
        
        with open(file_path, "wb") as buffer:
            buffer.write(content)
        
        # Processa IDs dos agentes
        agent_ids_list = []
        if agent_ids.strip():
            try:
                agent_ids_list = [int(id.strip()) for id in agent_ids.split(",") if id.strip()]
            except ValueError:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="IDs de agentes inválidos"
                )
        
        # Cria registro do conhecimento
        # Define status baseado no papel do usuário
        initial_status = models.KnowledgeStatus.APPROVED if current_user.role == models.UserRole.MASTER_ADMIN else models.KnowledgeStatus.PENDING
        
        new_knowledge = models.Knowledge(
            title=title,
            knowledge_type=models.KnowledgeType.DOCUMENT,
            file_path=file_path,
            file_type=file.content_type,
            tags=tags.strip() if tags.strip() else None,
            author_id=current_user.id,
            status=initial_status
        )
        
        # Se o Master Admin criou, já marca como aprovado
        if current_user.role == models.UserRole.MASTER_ADMIN:
            from datetime import datetime
            new_knowledge.approved_at = datetime.now()
            new_knowledge.approved_by_id = current_user.id
        
        db.add(new_knowledge)
        db.flush()
        
        # Processa o documento para extrair conteúdo (RAG)
        try:
            print(f"DEBUG: Iniciando processamento RAG do arquivo: {file.filename}")
            
            if file.content_type == "application/pdf":
                # Processa PDF
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(file_path)
                pages = loader.load_and_split()
                content = "\n".join([page.page_content for page in pages])
                print(f"DEBUG: PDF processado, {len(pages)} páginas, {len(content)} caracteres")
                
            elif file.content_type == "text/plain":
                # Processa TXT
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                print(f"DEBUG: TXT processado, {len(content)} caracteres")
                
            elif file.content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                # Processa DOCX/DOC
                try:
                    from langchain_community.document_loaders import Docx2txtLoader
                    loader = Docx2txtLoader(file_path)
                    documents = loader.load()
                    content = "\n".join([doc.page_content for doc in documents])
                    print(f"DEBUG: DOCX processado, {len(content)} caracteres")
                except ImportError:
                    # Fallback se docx2txt não estiver disponível
                    content = f"Documento carregado: {file.filename}"
                    print("DEBUG: DOCX processado com fallback")
            else:
                content = f"Documento carregado: {file.filename}"
                print("DEBUG: Tipo de arquivo não suportado para extração, usando fallback")
            
            # Verifica tamanho do conteúdo extraído
            MAX_TEXT_SIZE = 200000  # 200k caracteres (aproximadamente 200KB de texto)
            
            if len(content) > MAX_TEXT_SIZE:
                print(f"⚠️ AVISO: Documento muito grande ({len(content)} caracteres). Será truncado para {MAX_TEXT_SIZE} caracteres para otimizar o desempenho.")
                content = content[:MAX_TEXT_SIZE]
                content += "\n\n[DOCUMENTO TRUNCADO PARA OTIMIZAÇÃO - Upload documentos menores para melhor desempenho]"
            
            # Salva o conteúdo extraído no registro
            new_knowledge.content = content
            print(f"DEBUG: Conteúdo salvo no banco: {len(new_knowledge.content)} caracteres")
            
        except Exception as extract_error:
            print(f"ERRO na extração de conteúdo: {extract_error}")
            # Se falhar a extração, salva pelo menos o nome do arquivo
            new_knowledge.content = f"Documento: {file.filename} (erro na extração: {str(extract_error)})"
        
        # Associa com agentes
        if agent_ids_list:
            agents = db.query(models.Agent).filter(
                models.Agent.id.in_(agent_ids_list)
            ).all()
            new_knowledge.agents = agents
            print(f"DEBUG: Documento associado a {len(agents)} agentes")
        
        db.commit()
        db.refresh(new_knowledge)
        
        return schemas.KnowledgeUploadResponse(
            status="success",
            message=f"Documento '{file.filename}' adicionado à base de conhecimento",
            knowledge_id=new_knowledge.id,
            filename=file.filename
        )
        
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao processar arquivo: {str(e)}"
        )

# === ENDPOINTS DE INTERAÇÃO COM AGENTES (SISTEMA RAG) ===

@app.post("/agents/{agent_id}/upload", response_model=schemas.DocumentUploadResponse)
async def upload_document_for_agent(
    agent_id: int,
    file: UploadFile = File(...), 
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> schemas.DocumentUploadResponse:
    """
    Faz upload e processa um documento PDF para a base de conhecimento de um agente.
    
    Args:
        agent_id: ID do agente
        file: Arquivo PDF a ser processado
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        schemas.DocumentUploadResponse: Status do processamento
        
    Raises:
        HTTPException: Se o agente não for encontrado ou o arquivo não for PDF
    """
    # Verifica se o agente existe
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    # Verifica se o usuário é o dono do agente ou é admin
    if (db_agent.owner_id != current_user.id and 
        current_user.role not in [models.UserRole.ADMIN, models.UserRole.MASTER_ADMIN]):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions"
        )
    
    # Verifica se o arquivo é PDF, TXT ou DOCX
    allowed_types = ["application/pdf", "text/plain", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Apenas arquivos PDF, DOCX e TXT são permitidos."
        )
    
    # Cria diretórios necessários
    agent_chroma_path = f"{CHROMA_DB_PATH}/{agent_id}"
    if not os.path.exists(agent_chroma_path):
        os.makedirs(agent_chroma_path)
    
    if not os.path.exists(UPLOADS_PATH):
        os.makedirs(UPLOADS_PATH)
    
    # Inicializa o vector store do agente
    embedding_function = get_embeddings()
    if embedding_function is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Serviço de embeddings (Ollama) indisponível. Verifique se o Ollama está em execução e se o modelo está baixado."
        )

    agent_vectorstore = build_chroma(
        persist_directory=agent_chroma_path,
        embedding_function=embedding_function
    )
    
    # Salva temporariamente o arquivo
    temp_file_path = f"{UPLOADS_PATH}/{file.filename}"
    
    try:
        # Verifica se o arquivo tem conteúdo
        if not file.filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Nome do arquivo não fornecido"
            )
        
        # Lê o conteúdo do arquivo de forma segura
        try:
            content = await file.read()
            if not content:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Arquivo vazio ou não foi possível ler o conteúdo"
                )
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Erro ao ler o arquivo: {str(e)}"
            )
        
        # Salva o arquivo
        with open(temp_file_path, "wb") as buffer:
            buffer.write(content)
        
        # Verifica se o arquivo foi salvo corretamente
        if not os.path.exists(temp_file_path) or os.path.getsize(temp_file_path) == 0:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Erro ao salvar o arquivo temporariamente"
            )
        
        # Processa o arquivo baseado no tipo
        docs = []
        if file.content_type == "application/pdf":
            # Validação adicional para PDFs - verifica o cabeçalho do arquivo
            with open(temp_file_path, "rb") as f:
                header = f.read(4)
                if header != b'%PDF':
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="Arquivo não é um PDF válido. Verifique se o arquivo está correto."
                    )
            # Processa PDF
            try:
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(temp_file_path)
                docs = loader.load()
                if not docs:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="PDF não contém texto extraível ou está vazio"
                    )
                total_content = "".join([doc.page_content for doc in docs])
                if not total_content.strip():
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="PDF não contém texto válido para processamento"
                    )
            except HTTPException:
                raise
            except Exception as e:
                error_msg = str(e).lower()
                if "invalid pdf header" in error_msg:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="Arquivo não é um PDF válido. Verifique se o arquivo não está corrompido."
                    )
                elif "eof marker not found" in error_msg:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="PDF está incompleto ou corrompido. Tente novamente com um arquivo válido."
                    )
                elif "encrypted" in error_msg:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="PDF está protegido por senha. Use um PDF sem proteção."
                    )
                else:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Erro ao processar PDF: {str(e)}"
                    )
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Processa DOCX
            try:
                from docx import Document as DocxDocument
                from langchain.schema import Document
                docx_obj = DocxDocument(temp_file_path)
                full_text = "\n".join([para.text for para in docx_obj.paragraphs])
                if not full_text.strip():
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="DOCX não contém texto válido para processamento"
                    )
                docs = [Document(page_content=full_text, metadata={"source": file.filename, "type": "docx"})]
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Erro ao processar arquivo DOCX: {str(e)}"
                )
        else:
            # Processa arquivo de texto
            try:
                with open(temp_file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                if not content.strip():
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="Arquivo de texto está vazio"
                    )
            except UnicodeDecodeError:
                try:
                    with open(temp_file_path, "r", encoding="latin-1") as f:
                        content = f.read()
                except Exception as e:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Erro ao ler arquivo de texto: {str(e)}"
                    )
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Erro ao processar arquivo de texto: {str(e)}"
                )
            from langchain.schema import Document
            docs = [Document(page_content=content, metadata={"source": file.filename, "type": "text"})]
        
        # Divide o texto em chunks (otimizado para melhor recuperação)
        try:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=2000,  # Aumentado para capturar mais contexto
                chunk_overlap=400  # Aumentado proporcionalmente
            )
            splits = text_splitter.split_documents(docs)
            
            if not splits:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Não foi possível dividir o documento em partes"
                )
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Erro ao dividir documento: {str(e)}"
            )
        
        # Adiciona ao vector store
        try:
            agent_vectorstore.add_documents(documents=splits)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Erro ao adicionar documento ao banco vetorial: {str(e)}"
            )
        
        # Cria registro do documento no banco de dados
        try:
            document_record = models.Document(
                filename=file.filename,
                file_path=temp_file_path,
                file_type=file.content_type,
                agent_id=agent_id
            )
            
            db.add(document_record)
            db.commit()
            db.refresh(document_record)
        except Exception as e:
            db.rollback()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Erro ao salvar registro do documento: {str(e)}"
            )
        
        return schemas.DocumentUploadResponse(
            status="success",
            message=f"Documento '{file.filename}' processado para o agente {db_agent.name}.",
            filename=file.filename,
            document_id=document_record.id
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao processar o arquivo: {str(e)}"
        )
    finally:
        # Remove o arquivo temporário
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

# === FUNÇÕES AUXILIARES ===

def get_ai_model_config(db: Session) -> str:
    """
    Obtém a configuração atual do modelo de IA do sistema.
    
    Args:
        db: Sessão do banco de dados
        
    Returns:
        str: Tipo do modelo de IA (HYBRID, LLAMA_ONLY, GEMINI_ONLY)
    """
    config = db.query(models.SystemConfig).filter(
        models.SystemConfig.key == "ai_model_type"
    ).first()
    
    if config:
        return config.value
    
    # Configuração padrão
    return "HYBRID"

@app.post("/agents/{agent_id}/ask", response_model=schemas.AgentResponse)
async def ask_agent(
    agent_id: int,
    request: schemas.PromptRequest, 
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> schemas.AgentResponse:
    """
    Sistema de IA adaptativo com três modelos de operação configuráveis pelo Master Admin.
    
    HYBRID: Sistema completo em 3 estágios (padrão)
    - ESTÁGIO 1: RAG em Documentos com Llama
    - ESTÁGIO 2: RAG em Links com Gemini  
    - ESTÁGIO 3: Conhecimento Geral com Gemini
    
    LLAMA_ONLY: Apenas Llama local
    - Usa somente documentos carregados + Llama3
    - Não acessa links ou conhecimento geral
    
    GEMINI_ONLY: Apenas Gemini
    - Pula documentos locais, usa apenas Gemini
    - Pode usar links + conhecimento geral
    
    Args:
        agent_id: ID do agente
        request: Pergunta a ser feita
        db: Sessão do banco de dados
        current_user: Usuário autenticado
        
    Returns:
        schemas.AgentResponse: Resposta do agente com indicação do modelo usado
    """
    # Verifica se o agente existe e está aprovado
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent or db_agent.status != models.AgentStatus.APPROVED:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found or not approved"
        )
    
    # 🧪 MECANISMO DE TESTE - Forçar erro no Gemini para testar fallback Ollama
    FORCE_GEMINI_ERROR = False
    if "TESTE_FALLBACK_OLLAMA" in request.prompt.upper():
        FORCE_GEMINI_ERROR = True
        print("🧪 [TESTE] Forçando erro no Gemini para testar fallback Ollama")
        # Remove o comando de teste da pergunta
        request.prompt = request.prompt.replace("TESTE_FALLBACK_OLLAMA", "").replace("teste_fallback_ollama", "").strip()
    
    # Obtém configuração do modelo de IA
    ai_model_config = get_ai_model_config(db)
    print(f"[SISTEMA] Modo configurado: {ai_model_config}")
    print(f"[SISTEMA] Pergunta para agente {agent_id}: {request.prompt}")
    print(f"[SISTEMA] 🧪 Teste fallback ativo: {FORCE_GEMINI_ERROR}")
    
    # Prompt de sistema base
    system_prompt = db_agent.system_prompt or "Você é um assistente prestativo."
    
    try:
        # === MODO LLAMA_ONLY ===
        if ai_model_config == "LLAMA_ONLY":
            print("[MODO LLAMA] Usando apenas Llama local com documentos...")
            
            agent_chroma_path = f"{CHROMA_DB_PATH}/{agent_id}"
            if os.path.exists(agent_chroma_path):
                try:
                    embedding_function = get_embeddings()
                    llm_instance = get_llm()

                    if embedding_function is None or llm_instance is None:
                        raise HTTPException(
                            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                            detail="Serviço LLaMA indisponível. Verifique se o Ollama está rodando e se os modelos estão baixados."
                        )

                    agent_vectorstore = build_chroma(
                        persist_directory=agent_chroma_path,
                        embedding_function=embedding_function
                    )
                    retriever = agent_vectorstore.as_retriever(
                        search_kwargs={"k": 8}  # Aumenta número de documentos recuperados
                    )
                    relevant_docs = retriever.get_relevant_documents(request.prompt)
                    
                    if len(relevant_docs) > 0:
                        print(f"[MODO LLAMA] ✅ Encontrados {len(relevant_docs)} documentos relevantes")
                        
                        prompt_template_text = f"""
{system_prompt}

Use o contexto abaixo para responder à pergunta do usuário de forma precisa e detalhada. Se a resposta não for encontrada no contexto fornecido, responda exatamente: 'Eu não tenho informações sobre isso no meu conhecimento atual.'

Contexto:
{{context}}

Pergunta do usuário:
{{input}}

Resposta:"""
                        
                        prompt_template = ChatPromptTemplate.from_template(prompt_template_text)
                        document_chain = create_stuff_documents_chain(llm_instance, prompt_template)
                        retrieval_chain = create_retrieval_chain(retriever, document_chain)
                        
                        response = retrieval_chain.invoke({"input": request.prompt})
                        
                        print("[MODO LLAMA] ✅ Resposta processada com Llama!")
                        return schemas.AgentResponse(
                            response=response["answer"],
                            user=current_user.username,
                            note="Resposta com Llama local + documentos",
                            stage_used="Llama Only - Documentos"
                        )
                    else:
                        print("[MODO LLAMA] ❌ Nenhum documento relevante encontrado")
                        return schemas.AgentResponse(
                            response="Eu não tenho informações sobre isso no meu conhecimento atual.",
                            user=current_user.username,
                            note="Sem documentos relevantes no modo Llama",
                            stage_used="Llama Only - Sem documentos"
                        )
                        
                except Exception as e:
                    print(f"[MODO LLAMA] ❌ Erro: {e}")
                    return schemas.AgentResponse(
                        response="Erro ao processar sua pergunta.",
                        user=current_user.username,
                        note="Erro no processamento Llama",
                        stage_used="Llama Only - Erro"
                    )
            else:
                print("[MODO LLAMA] ❌ Nenhum documento carregado")
                return schemas.AgentResponse(
                    response="Este agente não possui documentos carregados.",
                    user=current_user.username,
                    note="Sem base de conhecimento no modo Llama",
                    stage_used="Llama Only - Sem base"
                )
        
        # === MODO GEMINI_ONLY ===
        elif ai_model_config == "GEMINI_ONLY":
            print("[MODO GEMINI] Usando apenas Gemini com base de conhecimento centralizada...")
            
            if not GOOGLE_API_KEY:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Google API Key não configurada"
                )
            
            # Busca conhecimentos associados ao agente que estão aprovados e não expirados
            print("[MODO GEMINI] Verificando base de conhecimento centralizada...")
            knowledge_query = db.query(models.Knowledge).join(
                models.knowledge_agent_association
            ).filter(
                models.knowledge_agent_association.c.agent_id == agent_id,
                models.Knowledge.status == models.KnowledgeStatus.APPROVED
            )
            
            # Filtra conhecimentos não expirados
            from datetime import datetime
            current_time = datetime.now()
            knowledge_items = knowledge_query.filter(
                (models.Knowledge.expires_at.is_(None)) | 
                (models.Knowledge.expires_at > current_time)
            ).all()
            
            if knowledge_items:
                print(f"[MODO GEMINI] ✅ Encontrados {len(knowledge_items)} conhecimentos ativos")
                
                # Coleta conteúdo de diferentes tipos de conhecimento
                knowledge_content = []
                
                for knowledge in knowledge_items:
                    content_text = ""
                    
                    if knowledge.knowledge_type == models.KnowledgeType.TEXT and knowledge.content:
                        content_text = knowledge.content
                        print(f"[MODO GEMINI] ✅ Usando conhecimento TEXT: {knowledge.title}")
                        
                    elif knowledge.knowledge_type == models.KnowledgeType.DOCUMENT and knowledge.content:
                        # Usa o conteúdo já extraído durante o upload
                        content_text = knowledge.content
                        print(f"[MODO GEMINI] ✅ Usando conhecimento DOCUMENT: {knowledge.title}")
                        
                    elif knowledge.knowledge_type == models.KnowledgeType.LINK and knowledge.url:
                        # Para links, usa o conteúdo armazenado ou faz scraping se necessário
                        if knowledge.content:
                            content_text = knowledge.content
                            print(f"[MODO GEMINI] ✅ Usando conhecimento LINK (cache): {knowledge.title}")
                        else:
                            try:
                                headers = {
                                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                                }
                                response = requests.get(knowledge.url, headers=headers, timeout=10)
                                response.raise_for_status()
                                
                                soup = BeautifulSoup(response.content, 'html.parser')
                                for script in soup(["script", "style", "nav", "footer", "header"]):
                                    script.decompose()
                                
                                text_content = soup.get_text()
                                lines = [line.strip() for line in text_content.splitlines()]
                                content_text = '\n'.join(line for line in lines if line)[:5000]
                                print(f"[MODO GEMINI] ✅ Conteúdo do link processado: {knowledge.url}")
                                
                            except Exception as e:
                                print(f"[MODO GEMINI] Erro ao processar link {knowledge.url}: {e}")
                                content_text = f"Link: {knowledge.url} (erro ao acessar)"
                    
                    if content_text:
                        # Limita o conteúdo para evitar timeout e melhorar performance
                        limited_content = content_text[:15000]  # Máximo 15k caracteres por conhecimento
                        if len(content_text) > 15000:
                            limited_content += "\n\n[Conteúdo truncado para otimização...]"
                        
                        knowledge_content.append({
                            "title": knowledge.title,
                            "type": knowledge.knowledge_type.value,
                            "content": limited_content,
                            "tags": knowledge.tags or ""
                        })
                
                if knowledge_content:
                    print(f"[MODO GEMINI] ✅ Preparando resposta com {len(knowledge_content)} fontes de conhecimento")
                    
                    # Calcula tamanho total do contexto
                    total_context_size = sum(len(item['content']) for item in knowledge_content)
                    print(f"[MODO GEMINI] 📊 Tamanho total do contexto: {total_context_size} caracteres")
                    
                    # Monta o contexto para o LLM com informação mais rica
                    context_text = "\n\n".join([
                        f"=== {item['type'].upper()}: {item['title']} ===\n"
                        f"Tags: {item['tags']}\n"
                        f"Conteúdo:\n{item['content']}"
                        for item in knowledge_content
                    ])
                    
                    gemini_prompt = f"""{system_prompt}

🎯 MISSÃO CRÍTICA: Você é um especialista em análise de documentos educacionais. Sua tarefa é encontrar e extrair QUALQUER informação relevante sobre o assunto perguntado.

📋 INSTRUÇÕES DETALHADAS:
1. 🔍 ANALISE METICULOSAMENTE todo o contexto fornecido linha por linha
2. 🎯 PROCURE por menções DIRETAS e INDIRETAS do assunto perguntado
3. 📝 CONSIDERE sinônimos, variações e referências relacionadas
4. ✅ Se encontrar QUALQUER informação relevante, responda com TODOS os detalhes encontrados
5. 📊 CITE trechos específicos e organize as informações de forma clara
6. ⚠️ APENAS responda "Eu não tenho informações sobre isso no meu conhecimento atual" se REALMENTE não existir NENHUMA informação relacionada

🔎 ESTRATÉGIA DE BUSCA:
- Procure pelo nome exato e variações
- Busque por palavras-chave relacionadas
- Identifique contextos e menções indiretas
- Analise tabelas, listas e seções estruturadas

Pergunta do usuário: {request.prompt}

📚 DOCUMENTO COMPLETO PARA ANÁLISE DETALHADA:
{context_text}

🎯 RESPOSTA DETALHADA (analise TODO o documento acima e extraia TODAS as informações encontradas):"""
                    
                    try:
                        # 🧪 Mecanismo de teste: forçar erro para testar fallback
                        if FORCE_GEMINI_ERROR:
                            print("🧪 [TESTE] Forçando erro no Gemini...")
                            raise Exception("Erro forçado para teste de fallback Ollama")
                            
                        model = genai.GenerativeModel(GEMINI_MODEL_NAME)
                        response = model.generate_content(gemini_prompt)
                        
                        # Verifica se encontrou resposta válida
                        if "não tenho informações sobre isso" not in response.text.lower():
                            print("[MODO GEMINI] ✅ Resposta encontrada na base de conhecimento!")
                            return schemas.AgentResponse(
                                response=response.text,
                                user=current_user.username,
                                note="Resposta baseada na base de conhecimento centralizada",
                                stage_used="Gemini Only - Base de Conhecimento"
                            )
                        else:
                            print("[MODO GEMINI] ❌ Informação não encontrada na base de conhecimento")
                            return schemas.AgentResponse(
                                response="Eu não tenho informações sobre isso no meu conhecimento atual.",
                                user=current_user.username,
                                note="Informação não encontrada na base de conhecimento",
                                stage_used="Gemini Only - Sem informação"
                            )
                            
                    except Exception as e:
                        print(f"[MODO GEMINI] ❌ Erro ao processar com Gemini: {e}")
                        
                        # Fallback para Ollama em caso de erro de quota ou outros problemas
                        print(f"[MODO GEMINI] 🔄 Tentando fallback para Ollama...")
                        try:
                            if not _ollama_available():
                                raise RuntimeError("Cliente Ollama não disponível")

                            chat_payload = {
                                'model': LLM_MODEL_NAME,
                                'messages': [{'role': 'user', 'content': gemini_prompt}],
                                'options': {
                                    'temperature': float(os.getenv("OLLAMA_TEMPERATURE", "0.2"))
                                }
                            }

                            client = get_ollama_client()
                            if client is not None:
                                ollama_response = client.chat(**chat_payload)
                            else:
                                ollama_response = ollama.chat(**chat_payload)
                            ollama_answer = ollama_response['message']['content']
                            print(f"[MODO GEMINI] ✅ Resposta obtida via Ollama: {len(ollama_answer)} caracteres")
                            
                            return schemas.AgentResponse(
                                response=ollama_answer,
                                user=current_user.username,
                                note="Resposta gerada via Ollama (fallback)",
                                stage_used="Gemini Only - Ollama Fallback"
                            )
                        except Exception as ollama_error:
                            print(f"[MODO GEMINI] ❌ Erro no fallback Ollama: {ollama_error}")
                            raise HTTPException(
                                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                                detail=f"Erro ao processar resposta: {str(e)} | Fallback Ollama: {str(ollama_error)}"
                            )
                else:
                    print("[MODO GEMINI] ❌ Nenhum conteúdo válido encontrado nos conhecimentos")
            else:
                print("[MODO GEMINI] ❌ Nenhum conhecimento ativo encontrado para este agente")
            
            # Se chegou até aqui, não encontrou na base de conhecimento
            print("[MODO GEMINI] Usando conhecimento geral...")
            gemini_prompt = f"""{system_prompt}

Pergunta: {request.prompt}

Resposta:"""
            
            try:
                model = genai.GenerativeModel(GEMINI_MODEL_NAME)
                response = model.generate_content(gemini_prompt)
                
                print("[MODO GEMINI] ✅ Resposta de conhecimento geral!")
                return schemas.AgentResponse(
                    response=response.text,
                    user=current_user.username,
                    note="Resposta baseada em conhecimento geral",
                    stage_used="Gemini Only - Conhecimento Geral"
                )
                
            except Exception as e:
                print(f"[MODO GEMINI] ❌ Erro ao gerar resposta: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Erro ao gerar resposta: {str(e)}"
                )
        
        # === MODO HYBRID (padrão) ===
        else:
            print("[MODO HÍBRIDO] Iniciando sistema de 3 estágios...")
            
            # === ESTÁGIO 1: RAG NA BASE DE CONHECIMENTO CENTRALIZADA ===
            print("[ESTÁGIO 1] Verificando base de conhecimento centralizada...")
            
            # Busca conhecimentos associados ao agente que estão aprovados e não expirados
            knowledge_query = db.query(models.Knowledge).join(
                models.knowledge_agent_association
            ).filter(
                models.knowledge_agent_association.c.agent_id == agent_id,
                models.Knowledge.status == models.KnowledgeStatus.APPROVED
            )
            
            # Filtra conhecimentos não expirados
            from datetime import datetime
            current_time = datetime.now()
            knowledge_items = knowledge_query.filter(
                (models.Knowledge.expires_at.is_(None)) | 
                (models.Knowledge.expires_at > current_time)
            ).all()
            
            if knowledge_items:
                print(f"[ESTÁGIO 1] ✅ Encontrados {len(knowledge_items)} conhecimentos ativos")
                
                # Coleta conteúdo de diferentes tipos de conhecimento
                knowledge_content = []
                
                for knowledge in knowledge_items:
                    content_text = ""
                    
                    if knowledge.knowledge_type == models.KnowledgeType.TEXT and knowledge.content:
                        content_text = knowledge.content
                        print(f"[ESTÁGIO 1] ✅ Usando conhecimento TEXT: {knowledge.title}")
                        
                    elif knowledge.knowledge_type == models.KnowledgeType.DOCUMENT and knowledge.content:
                        # Usa o conteúdo já extraído durante o upload
                        content_text = knowledge.content
                        print(f"[ESTÁGIO 1] ✅ Usando conhecimento DOCUMENT: {knowledge.title}")
                        
                    elif knowledge.knowledge_type == models.KnowledgeType.LINK and knowledge.url:
                        # Para links, faz scraping em tempo real
                        try:
                            headers = {
                                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                            }
                            response = requests.get(knowledge.url, headers=headers, timeout=10)
                            response.raise_for_status()
                            
                            soup = BeautifulSoup(response.content, 'html.parser')
                            for script in soup(["script", "style", "nav", "footer", "header"]):
                                script.decompose()
                            
                            text_content = soup.get_text()
                            lines = [line.strip() for line in text_content.splitlines()]
                            content_text = '\n'.join(line for line in lines if line)[:5000]  # Limita a 5k chars
                            print(f"[ESTÁGIO 1] ✅ Conteúdo do link processado: {knowledge.url}")
                            
                        except Exception as e:
                            print(f"[ESTÁGIO 1] Erro ao processar link {knowledge.url}: {e}")
                            content_text = f"Link: {knowledge.url} (erro ao acessar)"
                    
                    if content_text:
                        # NOVA ESTRATÉGIA: Busca vetorial inteligente para documentos grandes
                        max_content = 8000  # Limite conservador
                        
                        # Para documentos muito grandes (>50k), usar busca vetorial
                        if len(content_text) > 50000:
                            print(f"[ESTÁGIO 1] 📚 Documento grande detectado ({len(content_text)} chars), usando busca vetorial...")
                            
                            try:
                                # Configurar busca vetorial
                                embedding_function = get_embeddings()
                                if embedding_function is None:
                                    raise RuntimeError("Embeddings indisponíveis")
                                vectorstore = build_chroma(
                                    persist_directory=CHROMA_DB_PATH,
                                    embedding_function=embedding_function
                                )
                                
                                # Buscar chunks mais relevantes
                                docs = vectorstore.similarity_search(
                                    request.prompt,
                                    k=4,  # Apenas 4 chunks mais relevantes
                                    filter={"source": knowledge.title}
                                )
                                
                                if docs:
                                    relevant_content = "\n\n".join([doc.page_content for doc in docs])
                                    limited_content = f"BUSCA VETORIAL - Trechos mais relevantes para '{request.prompt}':\n\n{relevant_content}"
                                    print(f"[ESTÁGIO 1] ✅ Busca vetorial encontrou {len(docs)} chunks ({len(limited_content)} chars)")
                                else:
                                    print(f"[ESTÁGIO 1] ❌ Busca vetorial sem resultados, usando busca textual")
                                    raise Exception("Sem resultados vetoriais")
                                    
                            except Exception as e:
                                print(f"[ESTÁGIO 1] ⚠️ Erro na busca vetorial: {e}, usando busca textual otimizada")
                                # Fallback para busca textual super otimizada
                                limited_content = extract_smart_content(content_text, request.prompt, max_content)
                        else:
                            # Para documentos menores, usar método normal otimizado
                            limited_content = extract_smart_content(content_text, request.prompt, max_content)
                        
                        # Debug: verificar se o conteúdo limitado contém informações do Agrocolégio
                        if "agrocolégio" in limited_content.lower() or "agrocolegio" in limited_content.lower():
                            print(f"[ESTÁGIO 1] ✅ Conteúdo processado contém Agrocolégio!")
                        else:
                            print(f"[ESTÁGIO 1] ❌ Conteúdo processado NÃO contém Agrocolégio!")
                        
                        knowledge_content.append({
                            "title": knowledge.title,
                            "type": knowledge.knowledge_type.value,
                            "content": limited_content,
                            "tags": knowledge.tags or ""
                        })
                
                if knowledge_content:
                    print(f"[ESTÁGIO 1] ✅ Preparando resposta com {len(knowledge_content)} fontes de conhecimento")
                    
                    # Monta o contexto para o LLM
                    context_text = "\n\n".join([
                        f"[{item['type'].upper()}] {item['title']}\n"
                        f"Tags: {item['tags']}\n"
                        f"Conteúdo: {item['content']}"
                        for item in knowledge_content
                    ])
                    
                    # Debug: verificar se o contexto final contém informações do Agrocolégio
                    if "agrocolégio" in context_text.lower() or "agrocolegio" in context_text.lower():
                        print(f"[ESTÁGIO 1] ✅ CONTEXTO FINAL contém Agrocolégio!")
                        # Encontrar e mostrar trecho
                        content_lower = context_text.lower()
                        for term in ["agrocolégio", "agrocolegio"]:
                            pos = content_lower.find(term)
                            if pos != -1:
                                snippet = context_text[max(0, pos-50):pos+150]
                                print(f"[ESTÁGIO 1] 🎯 Trecho: ...{snippet}...")
                                break
                    else:
                        print(f"[ESTÁGIO 1] ❌ CONTEXTO FINAL NÃO contém Agrocolégio!")
                    
                    print(f"[ESTÁGIO 1] 📄 Preview do contexto enviado: {context_text[:500]}...")
                    print(f"[ESTÁGIO 1] � Tamanho total do contexto: {len(context_text)} caracteres")
                    print(f"[ESTÁGIO 1] �🔍 Pergunta: {request.prompt}")
                    
                    prompt_template_text = f"""
{system_prompt}

🎯 VOCÊ É UM ESPECIALISTA EM DOCUMENTOS EDUCACIONAIS DE GOIÁS

📋 MISSÃO: Encontrar e extrair TODAS as informações relevantes sobre: "{request.prompt}"

🔍 ESTRATÉGIA DE ANÁLISE OBRIGATÓRIA:
1. 📖 LEIA o documento COMPLETAMENTE, linha por linha
2. 🎯 PROCURE por menções DIRETAS do assunto (nome exato, siglas, referências)
3. 🔎 BUSQUE menções INDIRETAS (contexto, localização, atividades relacionadas)
4. 📊 ANALISE tabelas, listas, anexos e seções estruturadas
5. ✅ EXTRAIA TODOS os detalhes encontrados, por menores que sejam
6. 📝 ORGANIZE as informações de forma clara e detalhada

⚠️ REGRA CRÍTICA: 
- Se encontrar QUALQUER informação relacionada, responda com TODOS os detalhes
- CITE trechos específicos do documento
- NUNCA responda negativamente se houver informações no documento

Pergunta: {request.prompt}

📚 DOCUMENTO COMPLETO DE EDUCAÇÃO EM GOIÁS:
{context_text}

🎯 ANÁLISE COMPLETA E RESPOSTA DETALHADA:"""
                    
                    try:
                        # 🧪 Mecanismo de teste: forçar erro para testar fallback
                        if FORCE_GEMINI_ERROR:
                            print("🧪 [TESTE] Forçando erro no Gemini do Estágio 1...")
                            raise Exception("Erro forçado para teste de fallback Ollama no Estágio 1")
                            
                        model = genai.GenerativeModel(GEMINI_MODEL_NAME) if GOOGLE_API_KEY else None
                        if model:
                            response = model.generate_content(prompt_template_text)
                            
                            print(f"[ESTÁGIO 1] 📝 Resposta COMPLETA do Gemini: {response.text}")
                            
                            # Verifica se encontrou resposta válida (múltiplas variações)
                            resposta_lower = response.text.lower()
                            frases_negativas = [
                                "não tenho informações sobre isso",
                                "não tenho informações",
                                "não encontrei informações",
                                "não há informações",
                                "informações não disponíveis",
                                "não possuo informações"
                            ]
                            
                            frases_encontradas = [frase for frase in frases_negativas if frase in resposta_lower]
                            resposta_valida = not any(frase in resposta_lower for frase in frases_negativas)
                            
                            print(f"[ESTÁGIO 1] 🔍 Frases negativas encontradas: {frases_encontradas}")
                            print(f"[ESTÁGIO 1] 🔍 Resposta válida: {resposta_valida}")
                            print(f"[ESTÁGIO 1] 🔍 Tamanho da resposta: {len(response.text.strip())}")
                            
                            if resposta_valida and len(response.text.strip()) > 20:
                                print("[ESTÁGIO 1] ✅ Resposta encontrada na base de conhecimento!")
                                return schemas.AgentResponse(
                                    response=response.text,
                                    user=current_user.username,
                                    note="Resposta baseada na base de conhecimento centralizada",
                                    stage_used="1 - Base de Conhecimento Centralizada"
                                )
                            else:
                                print(f"[ESTÁGIO 1] ❌ Resposta não satisfatória ou negativa")
                        else:
                            print("[ESTÁGIO 1] ❌ Google API Key não configurada, pulando para próximo estágio")
                            
                    except Exception as e:
                        print(f"[ESTÁGIO 1] ❌ Erro ao processar com Gemini: {e}")
                        
                        # Fallback para Ollama quando Gemini falhar (quota ou outros erros)
                        print(f"[ESTÁGIO 1] 🔄 Tentando fallback para Ollama...")
                        try:
                            if not _ollama_available():
                                raise RuntimeError("Cliente Ollama não disponível")

                            chat_payload = {
                                'model': LLM_MODEL_NAME,
                                'messages': [{'role': 'user', 'content': prompt_template_text}],
                                'options': {
                                    'temperature': float(os.getenv("OLLAMA_TEMPERATURE", "0.2"))
                                }
                            }

                            client = get_ollama_client()
                            if client is not None:
                                ollama_response = client.chat(**chat_payload)
                            else:
                                ollama_response = ollama.chat(**chat_payload)
                            ollama_answer = ollama_response['message']['content']
                            print(f"[ESTÁGIO 1] ✅ Resposta obtida via Ollama: {len(ollama_answer)} caracteres")
                            
                            # Verificar se a resposta do Ollama é satisfatória
                            resposta_lower = ollama_answer.lower()
                            frases_negativas = [
                                "não tenho informações sobre isso",
                                "não tenho informações",
                                "não encontrei informações",
                                "não há informações",
                                "informações não disponíveis",
                                "não possuo informações"
                            ]
                            
                            resposta_valida = not any(frase in resposta_lower for frase in frases_negativas)
                            
                            if resposta_valida and len(ollama_answer.strip()) > 20:
                                print(f"[ESTÁGIO 1] ✅ Resposta válida do Ollama encontrada!")
                                return schemas.AgentResponse(
                                    response=ollama_answer,
                                    user=current_user.username,
                                    note="Resposta baseada na base de conhecimento via Ollama (fallback)",
                                    stage_used="1 - Base de Conhecimento (Ollama Fallback)"
                                )
                            else:
                                print(f"[ESTÁGIO 1] ❌ Resposta do Ollama também não satisfatória")
                                
                        except Exception as ollama_error:
                            print(f"[ESTÁGIO 1] ❌ Erro no fallback Ollama: {ollama_error}")
                else:
                    print("[ESTÁGIO 1] ❌ Nenhum conteúdo válido encontrado nos conhecimentos")
            else:
                print("[ESTÁGIO 1] ❌ Nenhum conhecimento ativo encontrado para este agente")
            
            # === ESTÁGIO 2: RAG EM LINKS COM GEMINI ===
            print("[ESTÁGIO 2] Verificando links salvos...")
            
            agent_links = db.query(models.Link).filter(models.Link.agent_id == agent_id).all()
            
            if agent_links and GOOGLE_API_KEY:
                print(f"[ESTÁGIO 2] Encontrados {len(agent_links)} links salvos")
                
                # Extrai conteúdo dos links por web scraping
                scraped_content = []
                for link in agent_links[:3]:  # Limita a 3 links para evitar timeout
                    try:
                        print(f"[ESTÁGIO 2] Fazendo scraping de: {link.url}")
                        
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                        }
                        response = requests.get(link.url, headers=headers, timeout=10)
                        response.raise_for_status()
                        
                        soup = BeautifulSoup(response.content, 'html.parser')
                        
                        # Remove elementos não relevantes
                        for script in soup(["script", "style", "nav", "footer", "header"]):
                            script.decompose()
                        
                        # Extrai texto
                        text_content = soup.get_text()
                        lines = [line.strip() for line in text_content.splitlines()]
                        content = '\n'.join(line for line in lines if line)
                        
                        if len(content) > 200:  # Só adiciona se tiver conteúdo suficiente
                            scraped_content.append({
                                "url": link.url,
                                "title": link.title or "Sem título",
                                "content": content[:2000]  # Limita o tamanho
                            })
                            
                    except Exception as e:
                        print(f"[ESTÁGIO 2] Erro ao fazer scraping de {link.url}: {e}")
                
                if scraped_content:
                    print(f"[ESTÁGIO 2] ✅ Conteúdo extraído de {len(scraped_content)} links")
                    print("[ESTÁGIO 2] Processando com Gemini...")
                    
                    # Prepara contexto com conteúdo dos links
                    context_text = "\n\n".join([
                        f"Link: {item['url']}\nTítulo: {item['title']}\nConteúdo: {item['content']}"
                        for item in scraped_content
                    ])
                    
                    gemini_prompt = f"""{system_prompt}

Use o contexto abaixo dos links salvos para responder à pergunta do usuário. Se a resposta não for encontrada no contexto fornecido, responda exatamente: 'Eu não tenho informações sobre isso no meu conhecimento atual.'

Contexto dos Links:
{context_text}

Pergunta: {request.prompt}

Resposta:"""
                    
                    model = genai.GenerativeModel(GEMINI_MODEL_NAME)
                    response = model.generate_content(gemini_prompt)
                    
                    # Verifica se encontrou resposta válida
                    if "não tenho informações sobre isso" not in response.text.lower():
                        print("[ESTÁGIO 2] ✅ Resposta encontrada nos links!")
                        return schemas.AgentResponse(
                            response=response.text,
                            user=current_user.username,
                            note="Resposta baseada nos links salvos",
                            stage_used="2 - RAG Links + Gemini"
                        )
            
            print("[ESTÁGIO 2] ❌ Nenhum link relevante ou erro no processamento")
            
            # === ESTÁGIO 3: CONHECIMENTO GERAL COM GEMINI ===
            print("[ESTÁGIO 3] Usando conhecimento geral do Gemini...")
            
            if not GOOGLE_API_KEY:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Google API Key não configurada"
                )
            
            gemini_prompt = f"{system_prompt}\n\nPergunta: {request.prompt}"
            
            model = genai.GenerativeModel(GEMINI_MODEL_NAME)
            response = model.generate_content(gemini_prompt)
            
            print("[ESTÁGIO 3] ✅ Resposta de conhecimento geral obtida!")
            return schemas.AgentResponse(
                response=response.text,
                user=current_user.username,
                note="Resposta de conhecimento geral",
                stage_used="3 - Conhecimento Geral Gemini"
            )
        
    except Exception as e:
        print(f"[SISTEMA] ❌ Erro crítico: {e}")
        
        # === FALLBACK ===
        return schemas.AgentResponse(
            response="Eu não tenho informações sobre isso no meu conhecimento atual.",
            user=current_user.username,
            note="Sistema em modo fallback devido a erro",
            stage_used="Fallback"
        )

# === ENDPOINT DE WEB SCRAPING ===

@app.post("/agents/{agent_id}/scrape-link", response_model=schemas.UrlScrapeResponse)
def scrape_and_add_to_knowledge(
    agent_id: int,
    url_data: schemas.UrlScrapeRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_user)
) -> Dict[str, Any]:
    """
    [ADMIN+] Extrai conteúdo de uma URL e adiciona à base de conhecimento do agente.
    
    Args:
        agent_id: ID do agente
        url_data: Dados da URL a ser processada
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões adequadas
        
    Returns:
        Dict: Status da operação e informações sobre o conteúdo extraído
        
    Raises:
        HTTPException: Se o agente não for encontrado ou houver erro no scraping
    """
    # Verificar se o agente existe
    db_agent = db.query(models.Agent).filter(models.Agent.id == agent_id).first()
    if not db_agent:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Agent not found"
        )
    
    # Verificar se o usuário tem permissão (dono do agente ou admin)
    if (db_agent.owner_id != current_user.id and 
        current_user.role not in [models.UserRole.ADMIN, models.UserRole.MASTER_ADMIN]):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions"
        )
    
    try:
        # Fazer requisição para a URL
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url_data.url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Extrair conteúdo com BeautifulSoup
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remover scripts, styles e outros elementos não relevantes
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Extrair título
        title = soup.find('title')
        title_text = title.get_text().strip() if title else "Sem título"
        
        # Extrair texto principal
        text_content = soup.get_text()
        
        # Limpar e normalizar o texto
        lines = [line.strip() for line in text_content.splitlines()]
        content = '\n'.join(line for line in lines if line)
        
        if not content or len(content.strip()) < 100:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Não foi possível extrair conteúdo suficiente da URL"
            )
        
        # Criar documento Langchain
        doc = Document(
            page_content=content,
            metadata={
                "source": url_data.url,
                "title": title_text,
                "type": "web_scraping",
                "agent_id": agent_id
            }
        )
        
        # Dividir o documento em chunks (otimizado para melhor recuperação)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,  # Aumentado para capturar mais contexto
            chunk_overlap=400  # Aumentado proporcionalmente
        )
        chunks = text_splitter.split_documents([doc])
        
        # Configurar embeddings
        embedding_function = get_embeddings()
        if embedding_function is None:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Serviço de embeddings (Ollama) indisponível."
            )

        # Adicionar ao Chroma (usar o mesmo padrão de diretório que o upload e ask)
        persist_directory = f"{CHROMA_DB_PATH}/{agent_id}"
        os.makedirs(persist_directory, exist_ok=True)
        print(f"[SCRAPE] Persistindo vetores em: {persist_directory}")
        vectorstore = build_chroma(
            persist_directory=persist_directory,
            embedding_function=embedding_function
        )
        
        # Adicionar documentos ao vectorstore
        vectorstore.add_documents(chunks)
        
        # Criar registro do link no banco de dados
        link_record = models.Link(
            url=url_data.url,
            title=title_text,
            description=f"Conteúdo extraído automaticamente via web scraping",
            agent_id=agent_id
        )
        
        db.add(link_record)
        db.commit()
        db.refresh(link_record)
        
        return schemas.UrlScrapeResponse(
            status="success",
            message=f"Conteúdo extraído e adicionado à base de conhecimento",
            url=url_data.url,
            title=title_text,
            content_length=len(content),
            chunks_processed=len(chunks)
        )
        
    except requests.RequestException as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Erro ao acessar a URL: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao processar o conteúdo: {str(e)}"
        )

# === ENDPOINTS DE GERENCIAMENTO DE USUÁRIOS ===

@app.patch("/master/users/{user_id}", response_model=schemas.User)
def update_user(
    user_id: int,
    user_update: schemas.UserUpdate,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> schemas.User:
    """
    [MASTER_ADMIN] Atualiza dados de um usuário.
    
    Args:
        user_id: ID do usuário a ser atualizado
        user_update: Dados para atualização
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Master Admin
        
    Returns:
        schemas.User: Usuário atualizado
        
    Raises:
        HTTPException: Se o usuário não for encontrado
    """
    # Não permite editar o próprio usuário para evitar lock-out
    if user_id == current_user.id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Você não pode editar o seu próprio usuário"
        )
    
    db_user = db.query(models.User).filter(models.User.id == user_id).first()
    if not db_user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Usuário não encontrado"
        )
    
    # Atualiza apenas os campos fornecidos
    update_data = user_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        if field == 'password':
            # Hash da nova senha
            setattr(db_user, 'hashed_password', auth.get_password_hash(value))
        else:
            setattr(db_user, field, value)
    
    db.commit()
    db.refresh(db_user)
    return db_user

@app.delete("/master/users/{user_id}")
def delete_user(
    user_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
):
    """
    [MASTER_ADMIN] Exclui um usuário do sistema.
    
    Args:
        user_id: ID do usuário a ser excluído
        db: Sessão do banco de dados
        current_user: Usuário autenticado com permissões de Master Admin
        
    Returns:
        dict: Mensagem de confirmação
        
    Raises:
        HTTPException: Se o usuário não for encontrado ou tentar excluir a si mesmo
    """
    # Não permite excluir o próprio usuário para evitar lock-out
    if user_id == current_user.id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Você não pode excluir o seu próprio usuário"
        )
    
    db_user = db.query(models.User).filter(models.User.id == user_id).first()
    if not db_user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Usuário não encontrado"
        )
    
    try:
        # Exclui o usuário
        db.delete(db_user)
        db.commit()
        
        return {"message": f"Usuário '{db_user.username}' excluído com sucesso"}
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao excluir usuário: {str(e)}"
        )

# === ENDPOINTS DE APROVAÇÃO DE CONHECIMENTO (MASTER ADMIN) ===

@app.post("/knowledge/{knowledge_id}/approve", response_model=schemas.KnowledgeApprovalResponse)
def approve_knowledge(
    knowledge_id: int,
    approval_request: schemas.KnowledgeApprovalRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(models.UserRole.MASTER_ADMIN))
) -> schemas.KnowledgeApprovalResponse:
    """
    [MASTER_ADMIN] Aprova ou rejeita um conhecimento pendente.
    
    Args:
        knowledge_id: ID do conhecimento
        approval_request: Dados da aprovação/rejeição
        db: Sessão do banco de dados
        current_user: Usuário Master Admin autenticado
        
    Returns:
        schemas.KnowledgeApprovalResponse: Resultado da aprovação
    """
    # Busca o conhecimento
    knowledge = db.query(models.Knowledge).filter(
        models.Knowledge.id == knowledge_id,
        models.Knowledge.status == models.KnowledgeStatus.PENDING
    ).first()
    
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conhecimento pendente não encontrado"
        )
    
    try:
        from datetime import datetime
        
        if approval_request.action == "approve":
            knowledge.status = models.KnowledgeStatus.APPROVED
            knowledge.approved_at = datetime.now()
            knowledge.approved_by_id = current_user.id
            knowledge.rejection_reason = None
            message = f"Conhecimento '{knowledge.title}' aprovado com sucesso"
            
        elif approval_request.action == "reject":
            knowledge.status = models.KnowledgeStatus.REJECTED
            knowledge.rejection_reason = approval_request.rejection_reason
            message = f"Conhecimento '{knowledge.title}' rejeitado"
            
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Ação inválida. Use 'approve' ou 'reject'"
            )
        
        db.commit()
        db.refresh(knowledge)
        
        return schemas.KnowledgeApprovalResponse(
            status="success",
            message=message,
            knowledge_id=knowledge.id
        )
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao processar aprovação: {str(e)}"
        )

# Execução principal
if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
